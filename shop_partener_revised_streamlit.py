"""
AI Shopping Assistant (Streamlit)

This module implements a Streamlit web app that:
- Accepts user shopping preferences via a form
- Uses an LLM-powered agent + web search tooling to fetch product recommendations
- Attempts to parse the agent response into a structured table
- Optionally exports the table as a Word document (.docx)

Environment variables
---------------------
- SERPER_API_KEY: required by `agno.tools.serper.SerperTools` for web search.
- OPENAI_API_KEY (or equivalent): required by `agno.models.openai.OpenAIChat`.
  The exact variable name depends on your OpenAI/Agno configuration.
"""

import streamlit as st
from dotenv import load_dotenv
import re
from io import BytesIO
from typing import Any, Dict, List, Mapping, Sequence

# Load environment variables (SERPER_API_KEY required)
load_dotenv()

PRODUCT_PATTERN = re.compile(
    r"\d+\.\s(.+?)\n"
    r"\s*-\s*Price:\s*₹([\d,\.]+)\n"
    r"\s*-\s*Fabric:\s*(.+?)\n"
    r"\s*-\s*Features:\s*(.+?)\n"
    r"\s*-\s*Link:\s*\[View on (.+?)\]\((https?://[^\)]+)\)",
    flags=re.MULTILINE,
)


@st.cache_resource
def setup_agent() -> Any:
    """
    Create (and cache) the recommendation agent.

    The agent uses:
    - `OpenAIChat(id="gpt-4o")` as its model
    - `SerperTools()` for real-time web search

    Streamlit caches this resource so repeated runs (e.g., widget updates) don't
    recreate the agent unnecessarily.

    Returns
    -------
    Any
        Configured Agno `Agent` instance.
    """
    # Import lazily to reduce initial app load time (only needed after user submits).
    from agno.agent import Agent
    from agno.models.openai import OpenAIChat
    from agno.tools.serper import SerperTools

    return Agent(
        name="shopping partner",
        model=OpenAIChat(id="gpt-4o"),
        instructions=[
            "You are a product recommender agent specializing in finding products that match user preferences.",
            "Prioritize finding products that satisfy as many user requirements as possible, but ensure a minimum match of 50%.",
            "Search for products only from authentic and trusted e-commerce websites such as Amazon, Flipkart, Myntra, Meesho, Google Shopping, Nike, and other reputable platforms.",
            "Verify that each product recommendation is in stock and available for purchase.",
            "Avoid suggesting counterfeit or unverified products.",
            "Clearly mention the key attributes of each product (e.g., price, brand, features) in the response.",
            "Format the recommendations neatly and ensure clarity for ease of user understanding.",
        ],
        tools=[SerperTools()],
        # The 'show_tool_calls' parameter was removed here
    )

# Function to extract structured info from response
def parse_response_to_rows(response_text: str) -> List[Dict[str, str]]:
    """
    Parse a recommendation text response into structured rows.

    This parser expects the agent response to follow a specific pattern for each
    product, for example:

        1. Product title
           - Price: ₹12,345
           - Fabric: Some fabric
           - Features: Some features
           - Link: [View on Amazon](https://example.com/item)

    Parameters
    ----------
    response_text:
        Full text produced by the agent (ideally after any markdown cleanup).

    Returns
    -------
    list[dict[str, str]]
        Rows with keys: Product, Price (₹), Fabric, Features, Store, Link.
        If no matches are found, the list will be empty.
    """
    matches = PRODUCT_PATTERN.findall(response_text)

    rows: List[Dict[str, str]] = []
    for match in matches:
        title, price, fabric, features, site, link = match
        rows.append({
            "Product": title.strip(),
            "Price (₹)": price.strip(),
            "Fabric": fabric.strip(),
            "Features": features.strip(),
            "Store": site.strip(),
            "Link": link.strip()
        })

    return rows

# DOCX export
@st.cache_data(show_spinner=False, max_entries=128)
def generate_docx_bytes_from_rows(
    rows: Sequence[Mapping[str, Any]],
    columns: Sequence[str],
) -> bytes:
    """
    Generate a Word document (.docx) containing the provided tabular data.

    The document includes a heading and a table (Table Grid style) containing
    the provided columns and values.

    Parameters
    ----------
    rows:
        Row mappings to export (e.g., list of dicts).
    columns:
        Column order to use in the document.

    Returns
    -------
    bytes
        DOCX bytes suitable for passing to Streamlit's `st.download_button`.
    """
    # Import lazily to reduce initial app load time (only needed when exporting).
    from docx import Document

    doc = Document()
    doc.add_heading('Product Recommendations', 0)
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    for i, column in enumerate(columns):
        hdr_cells[i].text = column

    for row in rows:
        cells = table.add_row().cells
        for i, column in enumerate(columns):
            cells[i].text = str(row.get(column, ""))

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def clean_markdown_bold(text: str) -> str:
    """Remove markdown bold markers (**...**) without changing content."""
    return re.sub(r"\*\*(.*?)\*\*", r"\1", text)


@st.cache_data(show_spinner=False, ttl=3600, max_entries=256)
def run_recommendation_query(query: str) -> str:
    """
    Run the (network-bound) recommendation query and cache results.

    Caching dramatically improves UX on Streamlit reruns (e.g., clicking Download)
    and reduces repeated LLM/search calls for identical queries.
    """
    agent = setup_agent()
    result = agent.run(query)
    return result.content if hasattr(result, "content") else str(result)

# --- UI ---
st.set_page_config(page_title="🛍️ AI Shopping Assistant", layout="centered")
st.title("🛍️ AI Shopping Assistant")
st.markdown("Find the best products that match your preferences using AI and real-time web search.")

if "recommendations_query" not in st.session_state:
    st.session_state.recommendations_query = ""
if "recommendations_text" not in st.session_state:
    st.session_state.recommendations_text = ""
if "recommendations_rows" not in st.session_state:
    st.session_state.recommendations_rows = []

with st.form("product_form"):
    col1, col2 = st.columns(2)
    with col1:
        category = st.text_input("Product Category", "Sports shoe")
    with col2:
        color = st.text_input("Preferred Color", "Blue")
    with col1:
        purpose = st.text_input("Purpose", "Comfortable for long-distance running")
    with col2:
        budget = st.number_input("Max Budget (INR)", min_value=0, value=10000, step=500)

    submitted = st.form_submit_button("🔍 Get Recommendations")

if submitted:
    with st.spinner("Searching and analyzing..."):
        query = (
            f"I am looking for {category} with the following preferences: "
            f"Color: {color}, Purpose: {purpose}, Budget: Under Rs. {budget}"
        )
        try:
            raw_text = run_recommendation_query(query)
        except Exception as e:
            st.error("Recommendation request failed. Please try again.")
            st.exception(e)
            st.stop()

        cleaned_text = clean_markdown_bold(raw_text)

        rows = parse_response_to_rows(cleaned_text)

        st.session_state.recommendations_query = query
        st.session_state.recommendations_text = cleaned_text
        st.session_state.recommendations_rows = rows

st.subheader("📋 Recommended Products")

rows = st.session_state.recommendations_rows
text = st.session_state.recommendations_text
if rows:
    st.dataframe(rows, use_container_width=True)

    columns = ["Product", "Price (₹)", "Fabric", "Features", "Store", "Link"]
    docx_bytes = generate_docx_bytes_from_rows(rows=rows, columns=columns)
    st.download_button(
        label="📥 Download as Word Document",
        data=docx_bytes,
        file_name="product_recommendations.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
elif text:
    # If parsing fails but the model produced useful prose, show it (persisted across reruns).
    if any(keyword in text.lower() for keyword in ["here are some", "options", "suitable"]):
        st.markdown(text)
    else:
        st.info("No product recommendations were found.")
